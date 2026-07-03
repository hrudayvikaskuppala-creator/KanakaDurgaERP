"""
===========================================
 Document Import Service
===========================================
Lets the user pick a PDF, Word (.docx) or image (.png/.jpg/.jpeg)
file - a supplier letterhead, a product datasheet, a purchase
invoice - and auto-fills a form with whatever this module can
confidently pull out of it.

Honesty note: this is heuristic/regex based, not a cloud AI parser
(the app runs fully offline). It gets straightforward, cleanly
formatted documents right most of the time and gets messy/handwritten
ones wrong sometimes - that's exactly why every import opens a review
screen instead of silently filling the database. Treat it as "auto-fill
+ verify", not "auto-fill + trust".

Supported inputs:
  - .pdf   -> text + tables via pdfplumber; falls back to OCR
              (pytesseract) on each page image if the PDF has no
              extractable text (i.e. it's a scanned document).
  - .docx  -> paragraphs + tables via python-docx.
  - .xlsx  -> every sheet's rows via openpyxl (each sheet becomes one
              table, exactly like a PDF/Word table - same downstream
              row-mapping code handles both).
  - .csv   -> rows via the standard library csv module (same table
              shape as above).
  - .png / .jpg / .jpeg -> OCR via pytesseract.

Requires: pdfplumber, python-docx, openpyxl, pytesseract, Pillow, and
the Tesseract-OCR binary installed on the machine (see requirements.txt).
"""

import csv
import os
import re


SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".xlsx", ".csv", ".png", ".jpg", ".jpeg")

# Files where a spreadsheet-like source (an actual table, not OCR/regex
# guesswork) can realistically hold *many* records - these are the
# formats offered first for bulk import (thousands of products/rows).
BULK_EXTENSIONS = (".xlsx", ".csv")


class ImportError_(Exception):
    pass


# ------------------------------------------------------------------
# Text / table extraction
# ------------------------------------------------------------------
def extract_text_and_tables(file_path):
    """
    Returns (text, tables) where tables is a list of tables, each a
    list of rows, each row a list of cell strings. Raises a friendly
    exception (with a clear message) if the file can't be read or a
    required library/binary is missing.

    Kept as a 2-tuple for backward compatibility with existing callers;
    see extract_text_and_tables_verbose() for the confidence flag too.
    """
    text, tables, _confidence = extract_text_and_tables_verbose(file_path)
    return text, tables


def extract_text_and_tables_verbose(file_path):
    """
    Same as extract_text_and_tables(), plus a "confidence" tag:
      - "table"    : came from a real table structure (xlsx/csv/pdf
                     table/docx table) - the most trustworthy source.
      - "text"     : came from a text layer with no table structure,
                     parsed with regex - reasonably trustworthy for
                     clean documents, less so for messy ones.
      - "ocr"      : came from OCR (scanned PDF or a photographed/
                     scanned image) - lowest confidence, always needs
                     a human check since handwriting, glare, skewed
                     photos etc. commonly cause misreads.
    Used to warn the person before they bulk-import OCR-sourced rows.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_from_pdf(file_path)
    elif ext == ".docx":
        text, tables = _extract_from_docx(file_path)
        return text, tables, ("table" if tables else "text")
    elif ext == ".xlsx":
        text, tables = _extract_from_xlsx(file_path)
        return text, tables, "table"
    elif ext == ".csv":
        text, tables = _extract_from_csv(file_path)
        return text, tables, "table"
    elif ext in (".png", ".jpg", ".jpeg"):
        return _extract_from_image(file_path), [], "ocr"
    else:
        raise ImportError_(
            f"Unsupported file type '{ext}'. Please choose a PDF, Word "
            f"(.docx), Excel (.xlsx), CSV, or image (.png/.jpg/.jpeg) file."
        )


def _extract_from_pdf(file_path):
    try:
        import pdfplumber
    except ImportError:
        raise ImportError_(
            "PDF import needs the 'pdfplumber' package. Install it with:\n"
            "pip install pdfplumber"
        )

    text_parts = []
    tables = []

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)

            for table in (page.extract_tables() or []):
                cleaned = [
                    [(cell or "").strip() for cell in row]
                    for row in table
                ]
                tables.append(cleaned)

        combined_text = "\n".join(text_parts).strip()
        confidence = "table" if tables else "text"

        # Scanned PDF (no extractable text layer) - fall back to OCR.
        if not combined_text:
            ocr_parts = []
            for page in pdf.pages:
                image = page.to_image(resolution=200).original
                ocr_parts.append(_ocr_image(image))
            combined_text = "\n".join(ocr_parts).strip()
            confidence = "ocr"

    return combined_text, tables, confidence


def _extract_from_docx(file_path):
    try:
        import docx
    except ImportError:
        raise ImportError_(
            "Word import needs the 'python-docx' package. Install it with:\n"
            "pip install python-docx"
        )

    document = docx.Document(file_path)

    text_parts = [p.text for p in document.paragraphs if p.text.strip()]

    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
        # Tables often repeat their text in paragraphs too; also add
        # them to the plain text so field regexes (GST, mobile, etc.)
        # can still find things that only appear inside a table cell.
        for row in rows:
            text_parts.append(" | ".join(row))

    return "\n".join(text_parts), tables


def _extract_from_xlsx(file_path):
    try:
        import openpyxl
    except ImportError:
        raise ImportError_(
            "Excel import needs the 'openpyxl' package. Install it with:\n"
            "pip install openpyxl"
        )

    try:
        workbook = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
    except Exception as exc:
        raise ImportError_(f"Couldn't open this Excel file.\n\nDetails: {exc}")

    tables = []
    text_parts = []

    for sheet in workbook.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            # Skip fully-blank rows (common at the end of exported sheets).
            if row is None or all(cell in (None, "") for cell in row):
                continue
            rows.append(["" if cell is None else str(cell).strip() for cell in row])

        if rows:
            tables.append(rows)
            text_parts.append(f"===== {sheet.title} =====")
            text_parts.extend(" | ".join(r) for r in rows)

    return "\n".join(text_parts), tables


def _extract_from_csv(file_path):
    # Excel commonly exports CSV with a UTF-8 BOM and/or a semicolon
    # delimiter depending on regional settings - handle both gracefully.
    try:
        with open(file_path, "r", newline="", encoding="utf-8-sig") as f:
            sample = f.read(4096)
            f.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample, delimiters=",;\t")
            except csv.Error:
                dialect = csv.excel
            reader = csv.reader(f, dialect)
            rows = [
                [cell.strip() for cell in row]
                for row in reader
                if any(cell.strip() for cell in row)
            ]
    except Exception as exc:
        raise ImportError_(f"Couldn't read this CSV file.\n\nDetails: {exc}")

    tables = [rows] if rows else []
    text_parts = [" | ".join(r) for r in rows]

    return "\n".join(text_parts), tables
    try:
        from PIL import Image
    except ImportError:
        raise ImportError_("Image import needs Pillow (already required by the app).")

    image = Image.open(file_path)
    return _ocr_image(image)


def _ocr_image(pil_image):
    try:
        import pytesseract
    except ImportError:
        raise ImportError_(
            "Image/scanned-PDF import needs the 'pytesseract' package "
            "(pip install pytesseract) plus the Tesseract-OCR program "
            "itself installed on this computer."
        )

    try:
        return pytesseract.image_to_string(pil_image)
    except Exception as exc:
        raise ImportError_(
            "Couldn't run OCR on this file. Make sure Tesseract-OCR is "
            f"installed and on your system PATH.\n\nDetails: {exc}"
        )


# ------------------------------------------------------------------
# Heuristic field parsing
# ------------------------------------------------------------------
_GSTIN_RE = re.compile(r"\b\d{2}[A-Z]{5}\d{4}[A-Z]\d[Z][A-Z\d]\b")
_MOBILE_RE = re.compile(r"(?<!\d)(?:\+?91[\s-]?)?([6-9]\d{9})(?!\d)")
_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")

# A bare "TAX INVOICE" heading must NOT match here - the "no/number/#/:"
# part is required so we only capture an actual identifier next to it.
_INVOICE_NO_RE = re.compile(
    r"(?:invoice|bill)\s*(?:no\.?|number|#|:)\s*[:\-]?\s*([A-Za-z0-9][A-Za-z0-9\-\/]{2,19})",
    re.IGNORECASE
)
_DATE_NEAR_LABEL_RE = re.compile(
    r"date\s*[:\-]?\s*(\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4})",
    re.IGNORECASE
)
_DATE_4DIGIT_YEAR_RE = re.compile(r"\b\d{1,2}[-/.]\d{1,2}[-/.]\d{4}\b")
_DATE_ANY_RE = re.compile(r"\b\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4}\b")
_PINCODE_RE = re.compile(r"\b\d{6}\b")

_COMPANY_KEYWORDS = (
    "pvt", "ltd", "enterprises", "batteries", "traders", "industries",
    "corporation", "company", "co.", "agencies", "distributors",
    "electricals", "motors", "automobiles"
)

_ADDRESS_KEYWORDS = (
    "road", "street", "nagar", "colony", "lane", "circle", "market",
    "complex", "building", "floor", "cross", "main", "district", "town"
)


def _clean_lines(text):
    return [line.strip() for line in text.splitlines() if line.strip()]


def parse_supplier_fields(text):
    """
    Best-effort extraction of supplier/company letterhead fields.
    Always returns every key (blank string if not found) so the
    review form can be built without extra None-checks.
    """
    lines = _clean_lines(text)

    result = {
        "supplier_name": "",
        "mobile": "",
        "gst_number": "",
        "email": "",
        "address": "",
    }

    gst_match = _GSTIN_RE.search(text)
    if gst_match:
        result["gst_number"] = gst_match.group(0)

    mobile_match = _MOBILE_RE.search(text)
    if mobile_match:
        result["mobile"] = mobile_match.group(1)

    email_match = _EMAIL_RE.search(text)
    if email_match:
        result["email"] = email_match.group(0)

    # Company name: prefer a line with a company-style keyword, else
    # fall back to the first substantial line in the document that
    # isn't obviously a document title.
    skip_words = ("invoice", "tax invoice", "quotation", "receipt", "bill of supply")
    name_candidate = ""

    for line in lines[:15]:
        lowered = line.lower()
        if any(k in lowered for k in _COMPANY_KEYWORDS):
            name_candidate = line
            break

    if not name_candidate:
        for line in lines[:5]:
            if line.lower() not in skip_words and len(line) > 3:
                name_candidate = line
                break

    result["supplier_name"] = name_candidate.strip(" '\"“”‘’•*-.")

    # Address: lines with address-style keywords or a 6-digit pincode,
    # joined together (usually 1-3 lines under the company name).
    address_lines = [
        line for line in lines
        if any(k in line.lower() for k in _ADDRESS_KEYWORDS) or _PINCODE_RE.search(line)
    ]
    result["address"] = ", ".join(address_lines[:3])

    return result


def parse_invoice_fields(text, tables=None):
    """
    Best-effort extraction for a purchase invoice: invoice number,
    date, supplier details, and product line items (from tables when
    available, otherwise a looser text-based guess).
    """
    tables = tables or []

    result = parse_supplier_fields(text)
    result["invoice_no"] = ""
    result["date"] = ""
    result["line_items"] = []

    invoice_match = _INVOICE_NO_RE.search(text)
    if invoice_match:
        result["invoice_no"] = invoice_match.group(1).strip(" :-")

    # Prefer a date right next to a "Date:" label (avoids matching
    # house-number-style addresses like "12-4-56, Bandar Road").
    # Then prefer a 4-digit year over a bare 2-digit one.
    date_match = (
        _DATE_NEAR_LABEL_RE.search(text)
        or _DATE_4DIGIT_YEAR_RE.search(text)
        or _DATE_ANY_RE.search(text)
    )
    if date_match:
        result["date"] = date_match.group(1) if date_match.re is _DATE_NEAR_LABEL_RE else date_match.group(0)

    result["line_items"] = _parse_line_items(text, tables)

    return result


def _parse_line_items(text, tables):
    items = _parse_line_items_from_tables(tables)
    if items:
        return items
    return _parse_line_items_from_text(text)


def _match_table_columns(header_row, header_aliases, required):
    """
    Given a table's header row and a {field: (alias, alias, ...)} map,
    returns {field: column_index} if at least all `required` fields
    were matched, else None. Shared by every "does this table look
    like an X table" check (line items, bulk products, bulk suppliers).
    """
    header = [cell.lower() for cell in header_row]
    column_index = {}

    for field, aliases in header_aliases.items():
        for i, cell in enumerate(header):
            if any(alias in cell for alias in aliases):
                column_index[field] = i
                break

    if not all(field in column_index for field in required):
        return None
    return column_index


def _row_number(row, column_index, field, default=0):
    idx = column_index.get(field)
    if idx is None or idx >= len(row):
        return default
    cleaned = re.sub(r"[^\d.]", "", row[idx] or "")
    try:
        return float(cleaned) if cleaned else default
    except ValueError:
        return default


def _row_text(row, column_index, field, default=""):
    idx = column_index.get(field)
    if idx is None or idx >= len(row):
        return default
    return (row[idx] or "").strip()


def _parse_line_items_from_tables(tables):
    """
    Looks for a table whose header row mentions product/qty/price-like
    columns, and maps every following row into a line item. This is
    far more reliable than text regex when the source document has an
    actual table (most invoices exported from other software do).
    """
    header_aliases = {
        "product": ("product", "item", "description", "particulars", "battery"),
        "qty": ("qty", "quantity", "nos"),
        "price": ("price", "rate", "unit price", "purchase price"),
        "gst": ("gst", "tax", "gst%", "tax%"),
        "total": ("total", "amount", "value"),
    }

    for table in tables:
        if len(table) < 2:
            continue

        header = [cell.lower() for cell in table[0]]
        column_index = {}

        for field, aliases in header_aliases.items():
            for i, cell in enumerate(header):
                if any(alias in cell for alias in aliases):
                    column_index[field] = i
                    break

        # Need at least a product name and one numeric column to be useful
        if "product" not in column_index or ("qty" not in column_index and "price" not in column_index):
            continue

        items = []
        for row in table[1:]:
            try:
                name = row[column_index["product"]].strip()
            except IndexError:
                continue

            if not name or name.lower() in ("total", "grand total", ""):
                continue

            def _num(field, default=0):
                idx = column_index.get(field)
                if idx is None or idx >= len(row):
                    return default
                cleaned = re.sub(r"[^\d.]", "", row[idx] or "")
                try:
                    return float(cleaned) if cleaned else default
                except ValueError:
                    return default

            items.append({
                "product": name,
                "qty": _num("qty", 1),
                "price": _num("price", 0),
                "gst": _num("gst", 0),
                "total": _num("total", 0),
            })

        if items:
            return items

    return []


_LINE_ITEM_TEXT_RE = re.compile(
    r"^(?P<name>[A-Za-z][A-Za-z0-9\-\s\/]{2,40}?)\s+"
    r"(?P<qty>\d{1,4})\s+"
    r"(?:[₹Rs\.\s]*)(?P<price>\d+(?:\.\d{1,2})?)\s+"
    r"(?:(?P<gst>\d{1,2}(?:\.\d)?)\s*%?\s+)?"
    r"(?:[₹Rs\.\s]*)(?P<total>\d+(?:\.\d{1,2})?)\s*$"
)


def _parse_line_items_from_text(text):
    """
    Fallback when no usable table was found: scans each line for a
    "name  qty  price  [gst%]  total" shape. Deliberately conservative
    - it's better to extract nothing (and let the user add items
    manually) than to guess wrong on unrelated numbers in the text.
    """
    items = []
    for line in _clean_lines(text):
        match = _LINE_ITEM_TEXT_RE.match(line)
        if not match:
            continue

        items.append({
            "product": match.group("name").strip(),
            "qty": float(match.group("qty")),
            "price": float(match.group("price")),
            "gst": float(match.group("gst")) if match.group("gst") else 0,
            "total": float(match.group("total")),
        })

    return items


PRODUCT_HEADER_ALIASES = {
    "brand": ("brand", "make"),
    "model": ("model", "battery model", "item code", "sku"),
    "capacity": ("capacity", "ah", "battery capacity"),
    "purchase_price": ("purchase price", "cost price", "buy price", "purchase rate"),
    "selling_price": ("selling price", "sale price", "mrp", "retail price", "sell rate"),
    "warranty": ("warranty",),
    "stock": ("stock", "quantity", "qty", "opening stock"),
    "rack_location": ("rack", "location", "rack location", "bin"),
}


def parse_products_from_table(table):
    """
    Maps a table (header row + data rows - from xlsx/csv/a PDF or docx
    table) into a list of product dicts using flexible header-name
    matching (e.g. "Sale Price", "MRP" and "Selling Price" all map to
    the same field). Returns [] if the table's headers don't look like
    a product list at all, so callers can safely try this on every
    table found in a file without false positives.

    Every row requires at minimum a brand or a model to be kept -
    fully blank rows and footer/total rows are silently skipped.
    """
    if len(table) < 2:
        return []

    column_index = _match_table_columns(
        table[0], PRODUCT_HEADER_ALIASES, required=("brand", "model")
    )
    if column_index is None:
        # Some supplier sheets only have a single combined "Product"
        # column (e.g. "Amaron 35Ah") rather than separate Brand/Model -
        # not something we can safely split automatically, so we bail
        # out to the text-based single-record fallback instead of
        # guessing and silently producing wrong brand/model pairs.
        return []

    items = []
    for row in table[1:]:
        brand = _row_text(row, column_index, "brand")
        model = _row_text(row, column_index, "model")

        if not brand and not model:
            continue
        if brand.lower() in ("total", "grand total") or model.lower() in ("total", "grand total"):
            continue

        items.append({
            "brand": brand,
            "model": model,
            "capacity": _row_text(row, column_index, "capacity"),
            "purchase_price": str(_row_number(row, column_index, "purchase_price")) or "",
            "selling_price": str(_row_number(row, column_index, "selling_price")) or "",
            "warranty": str(int(_row_number(row, column_index, "warranty"))) if column_index.get("warranty") is not None else "",
            "stock": str(int(_row_number(row, column_index, "stock"))) if column_index.get("stock") is not None else "",
            "rack_location": _row_text(row, column_index, "rack_location"),
        })

    return items


def parse_product_fields(text, tables=None):
    """
    Best-effort extraction for a single product datasheet/catalogue
    page. Prefers a real table when one is found in the file (far more
    reliable than guessing from free text) and only falls back to
    regex-on-text when no product-shaped table exists. Product
    datasheets vary wildly in layout, so this fills what it reasonably
    can and leaves the rest for the user - it's meant to save typing,
    not replace a final check.
    """
    tables = tables or []

    for table in tables:
        rows = parse_products_from_table(table)
        if rows:
            # Single-record callers (the PDF/Word "Import from File"
            # form) just want the first row; bulk callers should use
            # parse_products_from_table() directly instead.
            return rows[0]

    lines = _clean_lines(text)

    result = {
        "brand": "",
        "model": "",
        "capacity": "",
        "purchase_price": "",
        "selling_price": "",
        "warranty": "",
        "rack_location": "",
    }

    capacity_match = re.search(r"\b(\d{2,4})\s*Ah\b", text, re.IGNORECASE)
    if capacity_match:
        result["capacity"] = f"{capacity_match.group(1)}Ah"

    warranty_match = re.search(r"warranty[^\d]{0,15}(\d{1,3})\s*(month|year)", text, re.IGNORECASE)
    if warranty_match:
        months = int(warranty_match.group(1))
        if warranty_match.group(2).lower().startswith("year"):
            months *= 12
        result["warranty"] = str(months)

    price_matches = re.findall(r"(?:₹|Rs\.?)\s*([\d,]+(?:\.\d{1,2})?)", text)
    prices = []
    for raw in price_matches:
        try:
            prices.append(float(raw.replace(",", "")))
        except ValueError:
            pass
    if prices:
        prices.sort()
        result["purchase_price"] = str(prices[0])
        if len(prices) > 1:
            result["selling_price"] = str(prices[-1])

    model_match = re.search(r"\bmodel\s*[:\-]?\s*([A-Za-z0-9\-]+)", text, re.IGNORECASE)
    if model_match:
        result["model"] = model_match.group(1)

    # Brand: first line if it looks like a short product-title line
    # (battery brand names are usually 1-2 words at the top).
    for line in lines[:5]:
        words = line.split()
        if 1 <= len(words) <= 3 and line[0].isupper() and not any(ch.isdigit() for ch in line):
            result["brand"] = line
            break

    return result
